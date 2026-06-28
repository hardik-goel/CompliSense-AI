"""Agent-side artefact collector (extraction, local folder, classifier, staging)."""

import io
import json

from agent.collectors.base import Candidate, stage_candidates
from agent.collectors.classifier import Classification, classify, deterministic_classify
from agent.collectors.collect import collect_local
from agent.collectors.extract import extract_text
from agent.collectors.local_folder import crawl


# ---- classifier ----
def test_deterministic_matches_by_filename_and_keywords():
    c = deterministic_classify("privacy_notice.md", "We process personal data for the stated purpose of processing.")
    assert c.artefact_id == "privacy_notice" and c.confidence > 0.5 and c.method == "deterministic"


def test_deterministic_returns_none_for_unrelated():
    c = deterministic_classify("main.py", "import os\nprint('hello')")
    assert c.artefact_id is None and c.confidence == 0.0


def test_classify_falls_back_on_llm_error():
    class Boom:
        def available(self): return True
        def classify(self, *a): raise RuntimeError("api down")
    c = classify("privacy_notice.md", "personal data privacy policy", llm=Boom())
    assert c.artefact_id == "privacy_notice" and c.method == "deterministic"


def test_classify_uses_deterministic_when_llm_unavailable():
    class NoKey:
        def available(self): return False
        def classify(self, *a): raise AssertionError("should not be called")
    c = classify("consent_policy.md", "withdraw consent opt-in", llm=NoKey())
    assert c.artefact_id == "consent_policy" and c.method == "deterministic"


# ---- extraction ----
def test_extract_text_plain():
    assert "retention period" in extract_text(b"retention period 180 days", "retention.md")


def test_extract_pdf_roundtrip():
    import pypdf
    w = pypdf.PdfWriter(); w.add_blank_page(width=200, height=200)
    buf = io.BytesIO(); w.write(buf)
    assert isinstance(extract_text(buf.getvalue(), "doc.pdf"), str)  # must not raise


def test_extract_docx_roundtrip():
    import docx
    d = docx.Document(); d.add_paragraph("data retention schedule and erasure")
    buf = io.BytesIO(); d.save(buf)
    assert "retention schedule" in extract_text(buf.getvalue(), "policy.docx")


def test_extract_unknown_ext_empty():
    assert extract_text(b"\x89PNG", "logo.png") == ""


# ---- local crawl ----
def test_crawl_skips_binary_large_and_skip_dirs(tmp_path):
    (tmp_path / "privacy_notice.md").write_text("personal data, privacy policy")
    (tmp_path / "logo.png").write_bytes(b"\x89PNG\x00\x00")
    big = tmp_path / "huge.txt"; big.write_text("x" * (9 * 1024 * 1024))   # > 8MB cap
    skip = tmp_path / ".git"; skip.mkdir(); (skip / "config.yaml").write_text("retention period")
    refs = {c.ref for c in crawl(str(tmp_path))}
    assert "privacy_notice.md" in refs
    assert "logo.png" not in refs and "huge.txt" not in refs
    assert not any(".git" in r for r in refs)


def test_crawl_reads_docx(tmp_path):
    import docx
    d = docx.Document(); d.add_paragraph("encryption at rest, access control, least privilege")
    d.save(tmp_path / "security.docx")
    cands = crawl(str(tmp_path))
    assert any(c.filename == "security.docx" and "encryption" in c.sample for c in cands)


# ---- staging / collect ----
class _FakeLLM:
    def available(self): return True
    def classify(self, filename, sample):
        if "retention" in filename.lower():
            return Classification("retention_schedule", 0.9, "fake llm", "llm")
        return Classification(None, 0.1, "fake llm: none", "llm")


def test_collect_stages_matched_and_writes_manifest(tmp_path):
    src = tmp_path / "src"; src.mkdir()
    (src / "retention_schedule.md").write_text("retain personal data 180 days")
    (src / "random.txt").write_text("nothing relevant here")
    out = tmp_path / "out"
    m = collect_local(str(src), str(out), llm=_FakeLLM(), min_confidence=0.6)
    assert m["scanned"] == 2 and m["collected"] == 1
    assert (out / "retention_schedule.md").exists()
    manifest = json.loads((out / "COLLECTION_MANIFEST.json").read_text())
    assert manifest["matched"][0]["artefact_id"] == "retention_schedule"
    assert (out / "READ_ME_FIRST.txt").exists() and m["classifier"] == "llm"


def test_collect_deterministic_when_no_llm(tmp_path):
    src = tmp_path / "src"; src.mkdir()
    (src / "security_policy.md").write_text("encryption at rest, access control, least privilege")
    out = tmp_path / "out"
    m = collect_local(str(src), str(out), llm=None)
    assert m["collected"] == 1 and m["matched"][0]["artefact_id"] == "security_safeguards"
    assert m["classifier"] == "deterministic"


def test_stage_candidates_dedupes_names(tmp_path):
    cands = [
        Candidate(source="s3", ref="a/privacy_notice.md", filename="privacy_notice.md",
                  sample="personal data privacy policy purpose of processing", data=b"one"),
        Candidate(source="s3", ref="b/privacy_notice.md", filename="privacy_notice.md",
                  sample="personal data privacy policy purpose of processing", data=b"two"),
    ]
    out = tmp_path / "out"
    m = stage_candidates(cands, str(out), llm=None, source_label="s3")
    staged = sorted(x["staged_as"] for x in m["matched"])
    assert staged == ["privacy_notice.md", "privacy_notice_2.md"]
    assert (out / "privacy_notice.md").read_bytes() == b"one"
